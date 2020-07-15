from xlrd import open_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse

parser = argparse.ArgumentParser(description='Transform xls to doc')
parser.add_argument("--src")
args = parser.parse_args()
src = args.src

def copy_data(src_sheet, dest_doc):
    dest_doc.add_paragraph('\n' + src_sheet.name)

    table = dest_doc.add_table(rows=src_sheet.nrows + 1, cols=src_sheet.ncols + 3, style='Table Grid')
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'P1'
    hdr_cells[2].text = 'P2'
    hdr_cells[3].text = 'P3'
    hdr_cells[4].text = 'P4'
    hdr_cells[5].text = 'P5'
    hdr_cells[6].text = 'Примечание'
    
    for index_row in range(src_sheet.nrows):
        data_row = src_sheet.row_values(index_row)
        row_cells = table.rows[index_row+1].cells
        row_cells[0].text = str(index_row+1)
        row_cells[1].text = str(data_row[0])
        row_cells[2].text = str(data_row[1])
        row_cells[3].text = str(data_row[2])
        row_cells[4].text = ''
        row_cells[5].text = ''
        row_cells[6].text = str(data_row[3])

    return dest_doc

document = Document()
document.add_paragraph('Исходные данные').alignment = WD_ALIGN_PARAGRAPH.CENTER

read_book = open_workbook(src, on_demand=True)

for index_sheet in range(read_book.nsheets):
    read_sheet = read_book.get_sheet(index_sheet)
    document = copy_data(read_sheet, document)

document.save('demo.doc')
